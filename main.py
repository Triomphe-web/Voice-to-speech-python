import speech_recognition as sr
from googletrans import Translator
from docx import Document


# obtain audio from the microphone
def Listen():  # initializing funtion for voice to text
    r = sr.Recognizer()
    with sr.Microphone() as source:
        print("Dite le text a génerer !")
        # setting threshhold frequency for better output
        r.pause_threshhold = 1
        audio = r.listen(source, 0, 8)

    # recognize speech using google
    try:
        print("Recognition ... ")
        query = r.recognize_google(audio, language="fr")
    except sr.UnknownValueError:
        print("Could not understand audio")
        
    except sr.RequestError as e:
        print("Error Occured; {0}".format(e))
        
    except:
        return "text non reconnus"
    query = str(query).lower()  # converting the query into string
    return query


def printToWord():

    
    texteDite = Listen()

    document = Document()
    document.add_heading("Document génerée", level=1)

    document.add_paragraph(text=texteDite)

    document.save("document.docx")

    print("Text générer")




# printing the desired voice input into required text output
printToWord()